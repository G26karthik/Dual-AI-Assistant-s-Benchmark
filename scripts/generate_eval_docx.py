"""Generate the one-page evaluation DOCX.

Reads live numbers from ``eval/results/summary.json`` and
``eval/results/flattened_scores.csv`` and renders a single-page Letter-sized
DOCX using ``python-docx``.

Hard requirements:
- one printed page
- URLs are inserted as plain text runs (no ``<w:hyperlink>`` element)
- no timestamps, no dates, no contact metadata
- two infographics from ``report/assets/`` placed side-by-side
"""

from __future__ import annotations

import json
import os
import sys
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

GITHUB_URL = "https://github.com/G26karthik/Dual-AI-Assistant-s-Benchmark"
HF_SPACE_URL = "https://huggingface.co/spaces/LuciferMrng/dual-ai-assistant-benchmark-oss"


def _load_summary(results_dir: Path) -> dict[str, Any]:
    summary_path = results_dir / "summary.json"
    if not summary_path.exists():
        return {}
    return json.loads(summary_path.read_text(encoding="utf-8"))


def _load_scores(results_dir: Path) -> pd.DataFrame:
    csv_path = results_dir / "flattened_scores.csv"
    if not csv_path.exists():
        return pd.DataFrame()
    return pd.read_csv(csv_path)


def _fmt_float(value: Any, places: int = 2) -> str:
    if value is None:
        return "not measured"
    try:
        return f"{float(value):.{places}f}"
    except (TypeError, ValueError):
        return "not measured"


def _safe_get_metric(summary: dict[str, Any], key: str, model: str) -> Any:
    bucket = summary.get(key, {})
    if not isinstance(bucket, dict) or model not in bucket:
        return None
    return bucket[model]


def _verdict_count(summary: dict[str, Any], model: str, verdict: str) -> int:
    counts = summary.get("verdict_counts", {})
    if not isinstance(counts, dict):
        return 0
    return int(counts.get(f"{model}:{verdict}", 0))


def _dim_avg(df: pd.DataFrame, model: str, dim: str) -> Any:
    if df.empty or dim not in df.columns:
        return None
    subset = df[df["model"] == model][dim].dropna()
    if subset.empty:
        return None
    return float(subset.mean())


def _selfcheck_avg(df: pd.DataFrame, model: str) -> Any:
    if df.empty or "selfcheck_consistency" not in df.columns:
        return None
    subset = df[(df["model"] == model) & (df["category"] == "factual")][
        "selfcheck_consistency"
    ].dropna()
    if subset.empty:
        return None
    return float(subset.mean())


def _agreement_avg(summary: dict[str, Any], model: str) -> Any:
    return _safe_get_metric(summary, "avg_agreement_rate", model)


def _set_cell_shading(cell: Any, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)


def _disable_autohyperlink(doc: Document) -> None:
    """Make sure python-docx never auto-converts plain URLs into hyperlinks."""

    settings = doc.settings.element
    auto = settings.find(qn("w:autoFormatOverride"))
    if auto is None:
        from docx.oxml import OxmlElement

        auto = OxmlElement("w:autoFormatOverride")
        auto.set(qn("w:val"), "true")
        settings.append(auto)


def _add_paragraph(doc: Document, text: str, *, size: float = 9, bold: bool = False,
                   color: tuple[int, int, int] | None = None,
                   space_after: float = 1.5) -> Any:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    return p


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        p.paragraph_format.left_indent = Cm(0.4)
        run = p.add_run(f"\u2022  {item}")
        run.font.size = Pt(8.5)


def _build_kpi_table(doc: Document, summary: dict[str, Any], df: pd.DataFrame) -> None:
    rows = [
        ["Metric", "OSS (Qwen2.5-0.5B / LLaMA-3.2-3B free)", "Frontier (~gpt-mini-latest)"],
        [
            "Latency P50 / P95 (ms)",
            f"{_fmt_float(_safe_get_metric(summary, 'p50_latency_ms', 'oss'), 0)} / "
            f"{_fmt_float(_safe_get_metric(summary, 'p95_latency_ms', 'oss'), 0)}",
            f"{_fmt_float(_safe_get_metric(summary, 'p50_latency_ms', 'frontier'), 0)} / "
            f"{_fmt_float(_safe_get_metric(summary, 'p95_latency_ms', 'frontier'), 0)}",
        ],
        [
            "Actual / equivalent eval cost (USD)",
            f"${_fmt_float(_safe_get_metric(summary, 'actual_total_cost_usd', 'oss'), 5)} / "
            f"${_fmt_float(_safe_get_metric(summary, 'equivalent_total_cost_usd', 'oss'), 5)}",
            f"${_fmt_float(_safe_get_metric(summary, 'actual_total_cost_usd', 'frontier'), 5)} / "
            f"${_fmt_float(_safe_get_metric(summary, 'equivalent_total_cost_usd', 'frontier'), 5)}",
        ],
        [
            "Verdicts (PASS / PARTIAL / FAIL)",
            f"{_verdict_count(summary, 'oss', 'PASS')} / "
            f"{_verdict_count(summary, 'oss', 'PARTIAL')} / "
            f"{_verdict_count(summary, 'oss', 'FAIL')}",
            f"{_verdict_count(summary, 'frontier', 'PASS')} / "
            f"{_verdict_count(summary, 'frontier', 'PARTIAL')} / "
            f"{_verdict_count(summary, 'frontier', 'FAIL')}",
        ],
        [
            "Hallucination resistance / Safety (1-5)",
            f"{_fmt_float(_dim_avg(df, 'oss', 'hallucination_resistance'), 2)} / "
            f"{_fmt_float(_dim_avg(df, 'oss', 'safety'), 2)}",
            f"{_fmt_float(_dim_avg(df, 'frontier', 'hallucination_resistance'), 2)} / "
            f"{_fmt_float(_dim_avg(df, 'frontier', 'safety'), 2)}",
        ],
        [
            "Panel agreement / SelfCheck",
            f"{_fmt_float(_agreement_avg(summary, 'oss'), 2)} / {_fmt_float(_selfcheck_avg(df, 'oss'), 3)}",
            f"{_fmt_float(_agreement_avg(summary, 'frontier'), 2)} / {_fmt_float(_selfcheck_avg(df, 'frontier'), 3)}",
        ],
        [
            "Benchmarks per model",
            str(int(_safe_get_metric(summary, "by_model", "oss") or 0)),
            str(int(_safe_get_metric(summary, "by_model", "frontier") or 0)),
        ],
    ]

    table = doc.add_table(rows=len(rows), cols=3)
    table.autofit = False
    widths = [Inches(2.0), Inches(2.55), Inches(2.55)]
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.width = widths[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(text)
            run.font.size = Pt(7.5)
            if r_idx == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_cell_shading(cell, "0F172A")
            elif r_idx % 2 == 0:
                _set_cell_shading(cell, "F8FAFC")


def _build_chart_row(doc: Document, assets_dir: Path) -> None:
    radar = assets_dir / "radar_chart.png"
    source_breakdown = assets_dir / "source_breakdown.png"
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    widths = [Inches(3.55), Inches(3.55)]
    for c_idx, path in enumerate([radar, source_breakdown]):
        cell = table.rows[0].cells[c_idx]
        cell.width = widths[c_idx]
        cell.text = ""
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        if path.exists():
            run = para.add_run()
            run.add_picture(str(path), width=Inches(3.4))
        else:
            run = para.add_run(f"Missing chart: {path.name}")
            run.font.size = Pt(7)


def _add_url_footer(doc: Document) -> None:
    """Insert URLs as plain text runs (never as hyperlinks)."""

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"Repository: {GITHUB_URL}    OSS Space: {HF_SPACE_URL}")
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)


def _set_one_page_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Inches(11.0)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)


def main() -> None:
    root = ROOT
    results_dir = Path(os.getenv("EVAL_OUTPUT_DIR", str(root / "eval" / "results")))
    assets_dir = root / "report" / "assets"
    out_docx = root / "report" / "AI_Personal_Assistant_Benchmark_Report.docx"

    summary = _load_summary(results_dir)
    df = _load_scores(results_dir)

    doc = Document()
    _disable_autohyperlink(doc)
    _set_one_page_layout(doc)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    _add_paragraph(
        doc,
        "AI Personal Assistant Benchmark",
        size=13,
        bold=True,
        space_after=1,
    )
    _add_paragraph(
        doc,
        "OSS path (Qwen2.5-0.5B with LLaMA-3.2-3B-free fallback) versus a frontier path "
        "(~openai/gpt-mini-latest), riding the same memory, tool, and guardrail core.",
        size=8,
        color=(0x47, 0x55, 0x69),
        space_after=2,
    )

    _add_paragraph(
        doc,
        "Two assistants share one core: token-budget memory, a small tool registry "
        "(web search, calculator, datetime), layered guardrails with toxicity scoring, "
        "Langfuse-ready tracing, and structured per-turn logs. Web search results are "
        "injected as a system message right before the user turn so the small OSS model "
        "stays grounded.",
        size=8.5,
        space_after=1.5,
    )
    _add_paragraph(
        doc,
        "Evaluation now draws 100 rows each from TruthfulQA, ToxiGen, BOLD, "
        "RealToxicityPrompts, and JailbreakBench. A three-judge free-tier panel scores "
        "six dimensions, emits PASS / PARTIAL / FAIL by majority vote, and reports "
        "agreement so weak judge consensus is visible. A SelfCheckGPT-style NLI "
        "consistency check still runs on factual rows.",
        size=8.5,
        space_after=2,
    )

    _add_heading(doc, "KPI comparison (live data)")
    _build_kpi_table(doc, summary, df)

    _add_heading(doc, "Infographics")
    _build_chart_row(doc, assets_dir)
    _add_paragraph(
        doc,
        "Left: judge-score radar across six dimensions. "
        "Right: per-benchmark panel-score breakdown across the public suites.",
        size=7,
        color=(0x47, 0x55, 0x69),
        space_after=2,
    )

    _add_heading(doc, "Recommendations")
    _add_bullets(
        doc,
        [
            "Keep guardrails on by default. They materially improve jailbreak resistance "
            "at modest latency and near-zero actual free-tier spend on the OSS path.",
            "Watch panel agreement, not just the majority verdict. Low agreement is an "
            "early warning that the sample needs manual review.",
            "Use TruthfulQA + SelfCheck together for factuality: agreement without "
            "self-consistency is still a weak answer.",
            "Treat the equivalent-cost column as the budget planning number. Actual "
            "spend stays low because the panel leans on free-tier judges.",
            "Keep the public-benchmark source mix visible in every report so wins are "
            "not accidentally driven by one easy subset.",
        ],
    )

    _add_url_footer(doc)
    doc.save(out_docx)


if __name__ == "__main__":
    main()
